from docx import Document
import os

########################################################################################################################
# If you don't have bdfr (https://github.com/aliparlakci/bulk-downloader-for-reddit), install it with pip install bdfr #
########################################################################################################################

# First run the following command on the command line to download the stories as separate text files in a folder:
### py -m bdfr download C:\Users\your_user_name\your_folder\another_folder\reddit_TFTS\ --subreddit 'talesfromtechsupport' -S top -t all --file-scheme '{REDDITOR}_{TITLE}_{POSTID}_{DATE}' --make-hard-links

# Find all the txt files in the dataset folder
inputs = []
for file in os.listdir('C:\\Users\\your_user_name\\your_folder\\another_folder\\reddit_TFTS\\talesfromtechsupport'):
    if file.endswith(".txt"):
        inputs.append(os.path.join("C:\\Users\\your_user_name\\your_folder\\another_folder\\reddit_TFTS\\talesfromtechsupport", file))
 
 
# Concatenate all txt files in a file called my_written_file.docx and add a paragraph for every story
mydoc = Document()

mydoc.add_heading("https://www.reddit.com/r/talesfromtechsupport/ **Top stories**", 0)
mydoc.add_paragraph("\n")

for fname in inputs:
    with open(fname, encoding="utf-8", errors='ignore') as infile:
        mydoc.add_heading(str(fname[65:]), 1)
        mydoc.add_paragraph("\n")
        mydoc.add_paragraph(infile.read())
        mydoc.add_paragraph("\n")
mydoc.add_paragraph("\n\n\n----------END OF FILE----------")
mydoc.save("C:\\Users\\your_user_name\\your_folder\\another_folder\\reddit_TFTS\\my_written_file.docx")